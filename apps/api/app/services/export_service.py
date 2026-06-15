from __future__ import annotations

import io
import base64
from datetime import datetime


def export_to_pdf(title: str, content_native: str, content_target: str, keywords: list[str] = None, patterns: list[str] = None) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=20*mm, leftMargin=20*mm, topMargin=20*mm, bottomMargin=20*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CN', fontName='Helvetica', fontSize=11, leading=16))
    styles.add(ParagraphStyle(name='EN', fontName='Helvetica', fontSize=10, leading=14, textColor='#444444'))

    story = []
    story.append(Paragraph(f"<b>{title}</b>", styles['Title']))
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(f"<i>Exported on {datetime.now().strftime('%Y-%m-%d %H:%M')}</i>", styles['Normal']))
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph("<b>中文原文</b>", styles['Heading2']))
    for para in content_native.split('\n'):
        if para.strip():
            story.append(Paragraph(para, styles['CN']))
            story.append(Spacer(1, 2*mm))

    if content_target:
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph("<b>English Version</b>", styles['Heading2']))
        for para in content_target.split('\n'):
            if para.strip():
                story.append(Paragraph(para, styles['EN']))
                story.append(Spacer(1, 2*mm))

    if keywords:
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph("<b>Keywords</b>", styles['Heading2']))
        items = [ListItem(Paragraph(kw, styles['Normal'])) for kw in keywords[:20]]
        story.append(ListFlowable(items, bulletType='bullet'))

    if patterns:
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph("<b>Core Patterns</b>", styles['Heading2']))
        items = [ListItem(Paragraph(p, styles['EN'])) for p in patterns[:20]]
        story.append(ListFlowable(items, bulletType='bullet'))

    doc.build(story)
    return buf.getvalue()


def export_to_docx(title: str, content_native: str, content_target: str, keywords: list[str] = None, patterns: list[str] = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'

    doc.add_heading(title, 0)
    doc.add_paragraph(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("中文原文", level=1)
    for para in content_native.split('\n'):
        if para.strip():
            p = doc.add_paragraph(para)
            p.style.font.size = Pt(11)

    if content_target:
        doc.add_heading("English Version", level=1)
        for para in content_target.split('\n'):
            if para.strip():
                p = doc.add_paragraph(para)
                p.style.font.size = Pt(10)
                p.style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    if keywords:
        doc.add_heading("Keywords", level=2)
        for kw in keywords[:20]:
            doc.add_paragraph(kw, style='List Bullet')

    if patterns:
        doc.add_heading("Core Patterns", level=2)
        for p in patterns[:20]:
            doc.add_paragraph(p, style='List Bullet')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
